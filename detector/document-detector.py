import json
import uuid
import time
import boto3
import os
import re
import statistics
import logging
import io
import math
import requests
from collections import Counter, defaultdict
from urllib.parse import unquote_plus
import numpy as np

# PDF and DOCX processing libraries
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from docx import Document

# Configure logging for CloudWatch
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# --- Config ---
US_REGION = "us-east-1"
BEDROCK_REGION = os.environ.get("BEDROCK_REGION", US_REGION)
TABLE_NAME = os.environ.get("TABLE_NAME", "ContentDetections")
BUCKET_NAME = os.environ.get("BUCKET_NAME", "ai-detection-documents")

def get_cors_headers():
    """Return standardized CORS headers"""
    return {
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Headers": "Content-Type,X-Amz-Date,Authorization,X-Api-Key,X-Amz-Security-Token,Accept,Origin,Referer",
        "Access-Control-Allow-Methods": "GET,POST,OPTIONS,PUT,DELETE",
        "Access-Control-Max-Age": "86400",
        "Content-Type": "application/json"
    }

def handle_options_request():
    """Handle CORS preflight OPTIONS requests"""
    logger.info("Handling OPTIONS request for CORS preflight")
    return {
        "statusCode": 200,
        "headers": get_cors_headers(),
        "body": json.dumps({"message": "CORS preflight successful"})
    }

def get_clients():
    """Initialize AWS clients with forced region config."""
    data_region = os.environ.get("DATA_REGION", "ap-southeast-1")
    ai_region = os.environ.get("BEDROCK_REGION", "us-east-1")
    
    s3 = boto3.client("s3", region_name=data_region)
    dynamodb = boto3.resource("dynamodb", region_name=data_region)
    comprehend = boto3.client("comprehend", region_name=ai_region)
    bedrock = boto3.client("bedrock-runtime", region_name=ai_region)

    logger.info(f"S3 region: {s3.meta.region_name}")
    logger.info(f"DynamoDB region: {dynamodb.meta.client.meta.region_name}")
    logger.info(f"Comprehend region: {comprehend.meta.region_name}")
    logger.info(f"Bedrock region: {bedrock.meta.region_name}")

    return s3, dynamodb, comprehend, bedrock

def extract_text_from_pdf(file_stream):
    """Extract text from PDF using pdfminer.six."""
    try:
        text = pdf_extract_text(file_stream)
        
        if text and len(text.strip()) > 0:
            return text.strip()
        
        file_stream.seek(0)
        
        resource_manager = PDFResourceManager()
        fake_file_handle = io.StringIO()
        converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
        page_interpreter = PDFPageInterpreter(resource_manager, converter)
        
        for page in PDFPage.get_pages(file_stream, caching=True, check_extractable=True):
            page_interpreter.process_page(page)
        
        text = fake_file_handle.getvalue()
        converter.close()
        fake_file_handle.close()
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise

def extract_text_from_docx(file_stream):
    """Extract text from DOCX using python-docx."""
    try:
        doc = Document(file_stream)
        
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        extracted_text = '\n'.join(full_text)
        return extracted_text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise

def extract_text_from_document(s3, bucket_name, object_key):
    """Extract text from document using local libraries."""
    try:
        logger.info(f"Starting text extraction for {bucket_name}/{object_key}")
        
        file_extension = object_key.lower().split('.')[-1]
        
        if file_extension not in ['pdf', 'docx']:
            raise ValueError(f"Unsupported file type: {file_extension}. Only PDF and DOCX files are supported.")
        
        obj = s3.get_object(Bucket=bucket_name, Key=object_key)
        file_content = obj['Body'].read()
        file_stream = io.BytesIO(file_content)
        
        extracted_text = ""
        extraction_method = ""
        
        if file_extension == 'pdf':
            extracted_text = extract_text_from_pdf(file_stream)
            extraction_method = "pdfminer"
            logger.info(f"Extracted {len(extracted_text)} characters from PDF using pdfminer")
            
        elif file_extension == 'docx':
            extracted_text = extract_text_from_docx(file_stream)
            extraction_method = "python-docx"
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX using python-docx")
        
        if not extracted_text or len(extracted_text.strip()) < 10:
            raise ValueError(f"No readable text found in {file_extension.upper()} file or text too short")
        
        return extracted_text.strip(), extraction_method
        
    except Exception as e:
        logger.error(f"Error extracting text from document: {str(e)}")
        raise

# =============================================================================
# NEW RUBRIC-BASED SCORING SYSTEM
# =============================================================================

def analyze_perplexity_burstiness(text):
    """
    Category 1: Perplexity & Burstiness (15 pts)
    Analyzes sentence length variation, predictability, structural rhythm.
    """
    if not isinstance(text, str):
        return {"score": 7, "details": ["Invalid text input"]}
    
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
    
    if len(sentences) < 3:
        return {"score": 7, "details": "Too few sentences for analysis"}
    
    # Sentence length variation analysis
    sentence_lengths = [len(s.split()) for s in sentences]
    if len(sentence_lengths) > 1:
        mean_length = statistics.mean(sentence_lengths)
        std_length = statistics.stdev(sentence_lengths)
        cv = (std_length / mean_length) if mean_length > 0 else 0  # Coefficient of variation
    else:
        cv = 0
    
    # Sentence structure analysis
    complex_sentences = sum(1 for s in sentences if len(re.findall(r'[,;:]', s)) > 2)
    complex_ratio = complex_sentences / len(sentences)
    
    # Starting word patterns
    starters = [s.split()[0].lower() if s.split() else "" for s in sentences]
    starter_variety = len(set(starters)) / len(starters) if starters else 0
    
    # Scoring logic (0-15 points)
    score = 0
    details = []
    
    # High variation = more human (lower score)
    if cv > 0.6:  # High variation
        score += 2
        details.append("High sentence length variation (human-like)")
    elif cv > 0.4:  # Medium variation
        score += 5
        details.append("Medium sentence length variation")
    else:  # Low variation
        score += 10
        details.append("Low sentence length variation (AI-like)")
    
    # Structural predictability
    if complex_ratio > 0.6:  # Too uniform complexity
        score += 3
        details.append("Uniform complex structures (AI-like)")
    elif starter_variety < 0.5:  # Repetitive starters
        score += 2
        details.append("Repetitive sentence starters")
    
    return {
        "score": min(score, 15),
        "details": details,
        "coefficient_variation": cv,
        "complex_ratio": complex_ratio,
        "starter_variety": starter_variety
    }

def analyze_repetitiveness_redundancy(text):
    """
    Category 2: Repetitiveness & Redundancy (10 pts)
    Analyzes N-gram repetition, restating points, unnatural loops.
    """
    if not isinstance(text, str):
        return {"score": 5, "details": ["Invalid text input"]}
    
    words = re.findall(r'\b\w+\b', text.lower())
    
    if len(words) < 20:
        return {"score": 5, "details": "Text too short for analysis"}
    
    # N-gram analysis (3-5 word chunks)
    ngrams_3 = []
    ngrams_4 = []
    ngrams_5 = []
    
    for i in range(len(words) - 2):
        ngrams_3.append(' '.join(words[i:i+3]))
    
    for i in range(len(words) - 3):
        ngrams_4.append(' '.join(words[i:i+4]))
        
    for i in range(len(words) - 4):
        ngrams_5.append(' '.join(words[i:i+5]))
    
    # Count repetitions
    ngram3_counts = Counter(ngrams_3)
    ngram4_counts = Counter(ngrams_4)
    ngram5_counts = Counter(ngrams_5)
    
    repeated_3 = sum(1 for count in ngram3_counts.values() if count > 1)
    repeated_4 = sum(1 for count in ngram4_counts.values() if count > 1)
    repeated_5 = sum(1 for count in ngram5_counts.values() if count > 1)
    
    repetition_ratio = (repeated_3 + repeated_4 + repeated_5) / (len(ngrams_3) + len(ngrams_4) + len(ngrams_5)) if ngrams_3 else 0
    
    # Phrase repetition analysis
    sentences = re.split(r'[.!?]+', text)
    phrase_repetitions = 0
    for i, sentence in enumerate(sentences):
        for j, other_sentence in enumerate(sentences[i+1:], i+1):
            # Check for similar phrases
            sentence_words = set(sentence.lower().split())
            other_words = set(other_sentence.lower().split())
            if len(sentence_words) > 3 and len(other_words) > 3:
                overlap = len(sentence_words & other_words) / len(sentence_words | other_words)
                if overlap > 0.5:  # High overlap
                    phrase_repetitions += 1
    
    phrase_repetition_ratio = phrase_repetitions / len(sentences) if sentences else 0
    
    # Scoring (0-10 points)
    score = 0
    details = []
    
    if repetition_ratio > 0.15:  # High repetition
        score += 6
        details.append("High N-gram repetition (AI-like)")
    elif repetition_ratio > 0.08:
        score += 3
        details.append("Medium N-gram repetition")
    else:
        score += 1
        details.append("Low repetition (human-like)")
    
    if phrase_repetition_ratio > 0.2:
        score += 4
        details.append("High phrase repetition (AI-like)")
    elif phrase_repetition_ratio > 0.1:
        score += 2
        details.append("Some phrase repetition")
    
    return {
        "score": min(score, 10),
        "details": details,
        "repetition_ratio": repetition_ratio,
        "phrase_repetition_ratio": phrase_repetition_ratio
    }

def analyze_factuality_hallucination(text, bedrock):
    """
    Category 3: Factuality & Hallucination Detection (15 pts)
    Analyzes accuracy of claims, validity of facts, existence of cited entities.
    """
    try:
        # Use the existing fact-checking functionality
        fact_check_results = fact_check_document(text, bedrock)
        
        if fact_check_results and fact_check_results.get("status") == "success":
            factual_accuracy = fact_check_results.get("factual_accuracy", 50)
            hallucinations = len(fact_check_results.get("hallucinations_detected", []))
            misleading_claims = len(fact_check_results.get("misleading_claims", []))
            
            # Scoring logic (0-15 points)
            score = 0
            details = []
            
            # Convert factual accuracy to score (inverse relationship) - more sensitive
            if factual_accuracy > 85:  # High accuracy
                score += 1
                details.append("High factual accuracy (human-like)")
            elif factual_accuracy > 70:
                score += 3
                details.append("Good factual accuracy (human-like)")
            elif factual_accuracy > 50:
                score += 6
                details.append("Medium factual accuracy (concerning)")
            elif factual_accuracy > 30:
                score += 9
                details.append("Low factual accuracy (AI-like)")
            else:
                score += 12
                details.append("Very low factual accuracy (AI-like)")
            
            # Hallucination penalties
            if hallucinations > 3:
                score += 3
                details.append(f"Multiple hallucinations detected ({hallucinations})")
            elif hallucinations > 0:
                score += 1
                details.append(f"Some hallucinations detected ({hallucinations})")
            
            # Misleading claims penalties  
            if misleading_claims > 2:
                score += 2
                details.append(f"Multiple misleading claims ({misleading_claims})")
            elif misleading_claims > 0:
                score += 1
                details.append(f"Some misleading claims ({misleading_claims})")
            
            return {
                "score": min(score, 15),
                "details": details,
                "factual_accuracy": factual_accuracy,
                "hallucinations": hallucinations,
                "misleading_claims": misleading_claims
            }
        else:
            return {
                "score": 7,  # Neutral score when fact-checking fails
                "details": ["Fact-checking unavailable"],
                "factual_accuracy": 50,
                "hallucinations": 0,
                "misleading_claims": 0
            }
            
    except Exception as e:
        logger.warning(f"Fact-checking failed: {str(e)}")
        return {
            "score": 7,
            "details": ["Fact-checking failed"],
            "factual_accuracy": 50,
            "hallucinations": 0,
            "misleading_claims": 0
        }

def analyze_citation_reference_quality(text):
    """
    Category 4: Citation & Reference Quality (10 pts)
    Analyzes reference validity, correctness, and completeness.
    """
    if not isinstance(text, str):
        return {"score": 5, "details": ["Invalid text input"]}
    
    # Extract citations
    citations = extract_citations(text)
    citation_analysis = verify_citation_integrity(citations, text)
    
    total_citations = len(citations)
    suspicious_citations = len(citation_analysis.get("suspicious_citations", []))
    integrity_score = citation_analysis.get("integrity_score", 0)
    
    # Scoring logic (0-10 points)
    score = 0
    details = []
    
    if total_citations == 0:
        score += 3
        details.append("No citations found")
    else:
        # Based on citation integrity
        if integrity_score > 50:  # High suspicion
            score += 8
            details.append("High citation suspicion (AI-like)")
        elif integrity_score > 25:
            score += 5
            details.append("Medium citation concerns")
        else:
            score += 1
            details.append("Citations appear legitimate")
        
        # Suspicious citation ratio
        if total_citations > 0:
            suspicious_ratio = suspicious_citations / total_citations
            if suspicious_ratio > 0.3:
                score += 2
                details.append("High ratio of suspicious citations")
    
    return {
        "score": min(score, 10),
        "details": details,
        "total_citations": total_citations,
        "suspicious_citations": suspicious_citations,
        "integrity_score": integrity_score
    }

def analyze_stylistic_consistency(text):
    """
    Category 5: Stylistic Consistency (10 pts)
    Analyzes shifts in tone, presence of typos, uneven formality.
    """
    if not isinstance(text, str):
        return {"score": 5, "details": ["Invalid text input"]}
    
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
    
    if len(sentences) < 5:
        return {"score": 5, "details": "Too few sentences for analysis"}
    
    # Formality analysis
    formal_indicators = []
    informal_indicators = []
    
    for sentence in sentences:
        # Formal indicators
        formal_count = len(re.findall(r'\b(furthermore|moreover|consequently|therefore|thus|hence|nevertheless|nonetheless)\b', sentence.lower()))
        # Informal indicators  
        informal_count = len(re.findall(r'\b(gonna|wanna|kinda|sorta|yeah|ok|okay)\b', sentence.lower()))
        # Contractions
        contraction_count = len(re.findall(r"[a-z]'[a-z]", sentence.lower()))
        
        formal_indicators.append(formal_count)
        informal_indicators.append(informal_count + contraction_count)
    
    # Calculate consistency
    formal_variance = statistics.variance(formal_indicators) if len(formal_indicators) > 1 else 0
    informal_variance = statistics.variance(informal_indicators) if len(informal_indicators) > 1 else 0
    
    # Typo/error detection (very basic)
    potential_typos = len(re.findall(r'\b[a-z]{15,}\b', text))  # Very long words might be typos
    repeated_words = len(re.findall(r'\b(\w+)\s+\1\b', text, re.IGNORECASE))  # Repeated words
    
    # Punctuation consistency
    punct_density = len(re.findall(r'[.,;:!?]', text)) / len(text.split()) if text.split() else 0
    
    # Scoring logic (0-10 points)
    score = 0
    details = []
    
    # High consistency suggests AI
    total_variance = formal_variance + informal_variance
    if total_variance < 0.5:  # Very consistent
        score += 6
        details.append("Very consistent style (AI-like)")
    elif total_variance < 1.0:
        score += 3
        details.append("Moderately consistent style")
    else:
        score += 1
        details.append("Natural style variation (human-like)")
    
    # Natural human errors
    if potential_typos == 0 and repeated_words == 0:
        score += 2
        details.append("No natural errors detected (too perfect)")
    elif potential_typos > 0 or repeated_words > 0:
        score -= 1 if score > 0 else 0
        details.append("Natural errors present (human-like)")
    
    # Punctuation perfection
    if 0.12 <= punct_density <= 0.18:  # Very consistent punctuation
        score += 2
        details.append("Very consistent punctuation (AI-like)")
    
    return {
        "score": min(max(score, 0), 10),
        "details": details,
        "style_variance": total_variance,
        "potential_typos": potential_typos,
        "punctuation_density": punct_density
    }

def analyze_writing_style(text):
    """
    Category 6: Writing Style Analysis (10 pts)
    Analyzes voice, tone, grammar naturalness, personal expression.
    """
    if not isinstance(text, str):
        return {"score": 5, "details": ["Invalid text input"]}
    
    words = re.findall(r'\b\w+\b', text.lower())
    word_count = len(words)
    
    if word_count < 50:
        return {"score": 5, "details": "Text too short for analysis"}
    
    # Enhanced personal expression indicators
    personal_pronouns = len(re.findall(r'\b(i|me|my|myself|we|us|our|ourselves)\b', text.lower()))
    emotional_words = len(re.findall(r'\b(love|hate|angry|happy|sad|frustrated|excited|worried|scared|amazing|incredible|fantastic|terrible|awful|wonderful|disappointed|surprised|confused|curious)\b', text.lower()))
    personal_anecdotes = len(re.findall(r'\b(remember|recalled|experienced|happened|occurred|witnessed|felt|realized|discovered|noticed|observed|learned|understood)\b', text.lower()))
    
    # Voice uniqueness indicators
    idiomatic_expressions = len(re.findall(r'\b(piece of cake|break a leg|hit the hay|cost an arm and a leg|break the ice|piece of mind|in a nutshell|at the end of the day|once in a blue moon)\b', text.lower()))
    colloquialisms = len(re.findall(r'\b(yeah|yep|nope|gonna|wanna|kinda|sorta|pretty much|basically|actually|really|totally|absolutely)\b', text.lower()))
    
    # AI-specific mechanical indicators (enhanced)
    formal_connectors = len(re.findall(r'\b(furthermore|moreover|consequently|therefore|thus|hence|in conclusion|in summary|additionally|nevertheless|nonetheless)\b', text.lower()))
    generic_phrases = len(re.findall(r'\b(it is important to note|it should be noted|it is worth mentioning|in today\'s world|in modern society|it is widely known|it is commonly accepted|it is generally agreed)\b', text.lower()))
    
    # ChatGPT/LLM-specific patterns
    llm_patterns = len(re.findall(r'\b(let me explain|i\'ll help you|here\'s how|to get started|first, let\'s|now, let\'s|let\'s dive into|let\'s explore|let\'s take a look|let\'s examine)\b', text.lower()))
    academic_formality = len(re.findall(r'\b(it is crucial to|it is essential to|it is imperative that|it is vital to|it is necessary to|it is important to understand|it is worth noting)\b', text.lower()))
    
    # Calculate ratios
    personal_ratio = personal_pronouns / word_count if word_count > 0 else 0
    emotional_ratio = emotional_words / word_count if word_count > 0 else 0
    anecdote_ratio = personal_anecdotes / word_count if word_count > 0 else 0
    idiomatic_ratio = idiomatic_expressions / word_count if word_count > 0 else 0
    colloquial_ratio = colloquialisms / word_count if word_count > 0 else 0
    formal_ratio = formal_connectors / word_count if word_count > 0 else 0
    generic_ratio = generic_phrases / word_count if word_count > 0 else 0
    llm_ratio = llm_patterns / word_count if word_count > 0 else 0
    academic_ratio = academic_formality / word_count if word_count > 0 else 0
    
    # Enhanced scoring logic (0-10 points) - Higher score = more AI-like
    score = 0
    details = []
    
    # Personal expression scoring (more sensitive)
    if personal_ratio < 0.003:  # Very low personal pronouns (more sensitive threshold)
        score += 5
        details.append("Extremely low personal expression (AI-like)")
    elif personal_ratio < 0.008:
        score += 3
        details.append("Very low personal expression (AI-like)")
    elif personal_ratio < 0.015:
        score += 1
        details.append("Low personal expression")
    
    # Emotional content scoring (more sensitive)
    if emotional_ratio < 0.0005:  # Very low emotional content (more sensitive)
        score += 3
        details.append("Extremely low emotional content (AI-like)")
    elif emotional_ratio < 0.002:
        score += 2
        details.append("Very low emotional content (AI-like)")
    elif emotional_ratio < 0.005:
        score += 1
        details.append("Low emotional content")
    
    # Formal/mechanical language scoring (more sensitive)
    if formal_ratio > 0.008:  # High formal language (more sensitive)
        score += 3
        details.append("High formal language usage (AI-like)")
    elif formal_ratio > 0.004:
        score += 2
        details.append("Moderate formal language usage (AI-like)")
    elif formal_ratio > 0.002:
        score += 1
        details.append("Some formal language usage")
    
    # Generic phrases scoring (more sensitive)
    if generic_ratio > 0.003:  # High generic phrases (more sensitive)
        score += 3
        details.append("High generic phrase usage (AI-like)")
    elif generic_ratio > 0.001:
        score += 2
        details.append("Moderate generic phrase usage (AI-like)")
    elif generic_ratio > 0.0005:
        score += 1
        details.append("Some generic phrase usage")
    
    # LLM-specific pattern scoring (new)
    if llm_ratio > 0.002:  # LLM-specific patterns
        score += 4
        details.append("LLM-specific patterns detected (AI-like)")
    elif llm_ratio > 0.001:
        score += 2
        details.append("Some LLM-like patterns")
    
    # Academic formality scoring (new)
    if academic_ratio > 0.003:  # High academic formality
        score += 2
        details.append("High academic formality (AI-like)")
    elif academic_ratio > 0.001:
        score += 1
        details.append("Moderate academic formality")
    
    # Colloquial language bonus (human-like)
    if colloquial_ratio > 0.005:  # High colloquial usage
        score = max(0, score - 2)
        details.append("High colloquial usage (human-like)")
    elif colloquial_ratio > 0.002:
        score = max(0, score - 1)
        details.append("Some colloquial usage (human-like)")
    
    return {
        "score": min(score, 10),
        "details": details,
        "personal_ratio": personal_ratio,
        "emotional_ratio": emotional_ratio,
        "formal_ratio": formal_ratio,
        "generic_ratio": generic_ratio,
        "llm_ratio": llm_ratio,
        "academic_ratio": academic_ratio,
        "colloquial_ratio": colloquial_ratio
    }

def analyze_connectors_hedging(text):
    """
    Category 7: Overuse of Connectors & Hedging (10 pts)
    Analyzes "Additionally," "Moreover," "In summary," "It is important to note."
    """
    if not isinstance(text, str):
        return {"score": 5, "details": ["Invalid text input"]}
    
    word_count = len(text.split())
    
    if word_count < 50:
        return {"score": 5, "details": "Text too short for analysis"}
    
    # Connector words and phrases
    connectors = [
        "additionally", "moreover", "furthermore", "consequently", 
        "therefore", "thus", "hence", "nevertheless", "nonetheless",
        "however", "conversely", "alternatively", "similarly"
    ]
    
    # Hedging phrases
    hedging_phrases = [
        "it is important to note", "it should be noted", "it is worth mentioning",
        "it appears that", "it seems that", "it is likely that", "it is possible that",
        "in summary", "in conclusion", "to summarize", "overall"
    ]
    
    # Count occurrences
    connector_count = 0
    hedging_count = 0
    
    text_lower = text.lower()
    
    for connector in connectors:
        connector_count += len(re.findall(rf'\b{connector}\b', text_lower))
    
    for hedge in hedging_phrases:
        hedging_count += len(re.findall(rf'{hedge}', text_lower))
    
    # Calculate ratios
    connector_ratio = connector_count / word_count
    hedging_ratio = hedging_count / word_count
    total_ratio = connector_ratio + hedging_ratio
    
    # Scoring logic (0-10 points)
    score = 0
    details = []
    
    if total_ratio > 0.03:  # High usage
        score += 8
        details.append("High connector/hedging usage (AI-like)")
    elif total_ratio > 0.02:
        score += 5
        details.append("Medium connector/hedging usage")
    elif total_ratio > 0.01:
        score += 2
        details.append("Some connector/hedging usage")
    else:
        score += 1
        details.append("Natural connector usage (human-like)")
    
    # Specific penalties for AI-typical phrases
    if hedging_count > 2:
        score += 2
        details.append("Multiple hedging phrases detected")
    
    return {
        "score": min(score, 10),
        "details": details,
        "connector_count": connector_count,
        "hedging_count": hedging_count,
        "total_ratio": total_ratio
    }

def analyze_semantic_depth_originality(text):
    """
    Category 8: Semantic Depth & Originality (10 pts)
    Analyzes nuanced reasoning, domain-specific insight, creativity.
    """
    if not isinstance(text, str):
        return {"score": 5, "details": ["Invalid text input"]}
    
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
    
    if len(sentences) < 3:
        return {"score": 5, "details": "Too few sentences for analysis"}
    
    # Generic vs specific language
    generic_terms = len(re.findall(r'\b(various|numerous|several|many|different|important|significant|relevant|appropriate|effective|efficient|optimal|comprehensive|extensive|substantial)\b', text.lower()))
    specific_terms = len(re.findall(r'\b\d+\.?\d*\b|\b[A-Z][a-z]+\s[A-Z][a-z]+\b', text))  # Numbers and proper nouns
    
    # Shallow vs deep reasoning indicators
    shallow_phrases = len(re.findall(r'\b(in general|basically|simply put|obviously|clearly|evidently)\b', text.lower()))
    deep_reasoning = len(re.findall(r'\b(however|nevertheless|paradoxically|conversely|ironically|surprisingly|consequently|therefore|thus|hence)\b', text.lower()))
    
    # Creativity indicators
    metaphors_analogies = len(re.findall(r'\b(like|as if|similar to|resembles|analogous to|metaphorically|figuratively)\b', text.lower()))
    unique_perspectives = len(re.findall(r'\b(interestingly|remarkably|notably|curiously|unexpectedly|surprisingly)\b', text.lower()))
    
    word_count = len(text.split())
    
    # Calculate ratios
    generic_ratio = generic_terms / word_count if word_count > 0 else 0
    specific_ratio = specific_terms / word_count if word_count > 0 else 0
    shallow_ratio = shallow_phrases / word_count if word_count > 0 else 0
    
    # Scoring logic (0-10 points)
    score = 0
    details = []
    
    # High generic language suggests AI
    if generic_ratio > 0.02:
        score += 5
        details.append("High generic language usage (AI-like)")
    elif generic_ratio > 0.01:
        score += 2
        details.append("Some generic language")
    
    # Low specificity suggests AI
    if specific_ratio < 0.01:
        score += 3
        details.append("Low specificity (AI-like)")
    elif specific_ratio > 0.03:
        score -= 1 if score > 0 else 0
        details.append("Good specificity (human-like)")
    
    # Shallow reasoning suggests AI
    if shallow_ratio > 0.01:
        score += 2
        details.append("Shallow reasoning indicators")
    
    # Lack of creativity/originality
    if metaphors_analogies == 0 and unique_perspectives == 0:
        score += 2
        details.append("Limited creativity indicators")
    elif metaphors_analogies > 0 or unique_perspectives > 0:
        score -= 1 if score > 0 else 0
        details.append("Some creative elements present")
    
    return {
        "score": min(max(score, 0), 10),
        "details": details,
        "generic_ratio": generic_ratio,
        "specific_ratio": specific_ratio,
        "shallow_ratio": shallow_ratio
    }

def analyze_metadata_time_clues(text, metadata=None):
    """
    Category 9: Metadata & Time-Based Clues (10 pts)
    Analyzes document creation time, structure similarity across submissions.
    Note: Limited metadata available in this context.
    """
    if not isinstance(text, str):
        return {"score": 5, "details": ["Invalid text input"]}
    
    # Time references in text
    temporal_references = len(re.findall(r'\b(today|yesterday|tomorrow|currently|recently|lately|now|then|earlier|later|before|after|during|while|when)\b', text.lower()))
    
    # Specific dates and times
    specific_dates = len(re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}\b|\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2}\b', text.lower()))
    
    # Document structure indicators
    sections = len(re.findall(r'^[A-Z][a-z]+:|\n[A-Z][A-Z\s]+\n|\d+\.\s+[A-Z]', text, re.MULTILINE))
    
    # AI-typical perfect structure
    perfect_structure_indicators = 0
    if re.search(r'Introduction.*Methodology.*Results.*Discussion.*Conclusion', text, re.IGNORECASE | re.DOTALL):
        perfect_structure_indicators += 1
    if len(re.findall(r'^\d+\.\s+', text, re.MULTILINE)) > 5:  # Many numbered points
        perfect_structure_indicators += 1
    
    word_count = len(text.split())
    temporal_ratio = temporal_references / word_count if word_count > 0 else 0
    
    # Scoring logic (0-10 points)
    score = 0
    details = []
    
    # Low temporal references suggest AI (lacks grounding in time)
    if temporal_ratio < 0.01:
        score += 4
        details.append("Few temporal references (AI-like)")
    elif temporal_ratio < 0.02:
        score += 2
        details.append("Some temporal references")
    else:
        score += 1
        details.append("Good temporal grounding (human-like)")
    
    # Overly perfect structure suggests AI
    if perfect_structure_indicators > 1:
        score += 4
        details.append("Overly perfect structure (AI-like)")
    elif perfect_structure_indicators == 1:
        score += 2
        details.append("Some structural perfection")
    
    # Lack of specific dates/times
    if specific_dates == 0 and word_count > 200:
        score += 2
        details.append("No specific dates/times mentioned")
    
    return {
        "score": min(score, 10),
        "details": details,
        "temporal_ratio": temporal_ratio,
        "specific_dates": specific_dates,
        "perfect_structure_indicators": perfect_structure_indicators
    }

def calculate_rubric_score(text, bedrock, metadata=None):
    """
    Calculate the overall AI detection score using the 9-category rubric.
    Total: 100 points
    Thresholds: 0-40 (Likely Human), 41-70 (Suspicious/Mixed), 71-100 (Likely AI)
    """
    try:
        # Category 1: Perplexity & Burstiness (15 pts)
        category1 = analyze_perplexity_burstiness(text)
        
        # Category 2: Repetitiveness & Redundancy (10 pts)
        category2 = analyze_repetitiveness_redundancy(text)
        
        # Category 3: Factuality & Hallucination Detection (15 pts)
        category3 = analyze_factuality_hallucination(text, bedrock)
        
        # Category 4: Stylistic Consistency (10 pts)
        category4 = analyze_stylistic_consistency(text)
        
        # Category 5: Writing Style Analysis (10 pts)
        category5 = analyze_writing_style(text)
        
        # Category 6: Overuse of Connectors & Hedging (10 pts)
        category6 = analyze_connectors_hedging(text)
        
        # Category 7: Semantic Depth & Originality (10 pts)
        category7 = analyze_semantic_depth_originality(text)
        
        # Category 8: Metadata & Time-Based Clues (10 pts)
        category8 = analyze_metadata_time_clues(text, metadata)
        
        # LLM-specific pattern detection (bonus scoring)
        llm_analysis = detect_llm_patterns(text)
        llm_bonus = min(llm_analysis["llm_score"] // 5, 10)  # Convert to 0-10 scale
        
        # Calculate total score with LLM bonus (excluding citation quality)
        total_score = (
            category1["score"] +
            category2["score"] + 
            category3["score"] +
            category4["score"] +
            category5["score"] +
            category6["score"] +
            category7["score"] +
            category8["score"] +
            llm_bonus
        )
        
        # Determine classification with more sensitive thresholds
        if total_score <= 25:
            classification = "Likely Human"
            confidence = "high" if total_score <= 15 else "medium"
        elif total_score <= 50:
            classification = "Suspicious/Mixed"
            confidence = "medium"
        else:
            classification = "Likely AI"
            confidence = "high" if total_score >= 70 else "medium"
        
        # Generate detailed explanation
        explanation = generate_rubric_explanation(
            total_score, classification, confidence,
            [category1, category2, category3, category4, category5, 
             category6, category7, category8]
        )
        
        return {
            "total_score": total_score,
            "classification": classification,
            "confidence": confidence,
            "explanation": explanation,
            "category_breakdown": {
                "perplexity_burstiness": category1,
                "repetitiveness_redundancy": category2,
                "factuality_hallucination": category3,
                "stylistic_consistency": category4,
                "writing_style": category5,
                "connectors_hedging": category6,
                "semantic_depth": category7,
                "metadata_time_clues": category8,
                "llm_patterns": {
                    "score": llm_bonus,
                    "details": llm_analysis["detected_patterns"],
                    "pattern_count": llm_analysis["pattern_count"],
                    "confidence": llm_analysis["confidence"]
                }
            }
        }
        
    except Exception as e:
        logger.error(f"Rubric scoring error: {str(e)}")
        return {
            "total_score": 50,
            "classification": "Unknown",
            "confidence": "low",
            "explanation": f"Analysis error: {str(e)}",
            "category_breakdown": {}
        }

def generate_rubric_explanation(total_score, classification, confidence, categories):
    """Generate concise explanation based on rubric analysis (max 150 words)."""
    
    # Start with core information
    explanation_parts = [f"Score: {total_score}/90 → {classification} ({confidence} confidence)"]
    
    # Find top 3 most significant categories (excluding citation quality)
    category_names = [
        "Perplexity & Burstiness", "Repetitiveness & Redundancy", "Factuality & Hallucination",
        "Stylistic Consistency", "Writing Style", "Connectors & Hedging", 
        "Semantic Depth", "Metadata & Time Clues"
    ]
    max_scores = [15, 10, 15, 10, 10, 10, 10, 10]
    
    # Get top 3 categories by score significance
    significant_categories = []
    for i, category in enumerate(categories):
        if isinstance(category, dict) and "score" in category:
            score = category["score"]
            max_score = max_scores[i]
            significance = score / max_score if max_score > 0 else 0
            significant_categories.append((i, category, significance, score, max_score))
    
    # Sort by significance and take top 3
    significant_categories.sort(key=lambda x: x[2], reverse=True)
    top_categories = significant_categories[:3]
    
    # Add top categories with brief details
    for i, category, significance, score, max_score in top_categories:
        details = category.get("details", [])
        main_detail = details[0] if details else "Standard analysis"
        explanation_parts.append(f"{category_names[i]}: {score}/{max_score} - {main_detail}")
    
    # Add key indicators based on score
    if total_score >= 71:
        explanation_parts.append("Key AI indicators detected in multiple categories.")
    elif total_score <= 40:
        explanation_parts.append("Strong human writing patterns identified.")
    else:
        explanation_parts.append("Mixed content with both human and AI characteristics.")
    
    # Join and limit to 150 words
    full_explanation = " ".join(explanation_parts)
    words = full_explanation.split()
    if len(words) > 150:
        full_explanation = " ".join(words[:150]) + "..."
    
    return full_explanation

# =============================================================================
# LLM-SPECIFIC DETECTION PATTERNS
# =============================================================================

def detect_llm_patterns(text):
    """
    Detect specific patterns from popular LLMs like ChatGPT, Claude, DeepSeek, etc.
    Returns a score indicating likelihood of LLM generation.
    """
    llm_score = 0
    detected_patterns = []
    
    # ChatGPT-specific patterns
    chatgpt_patterns = [
        r'\b(Python is a high-level programming language|Let me explain|I\'ll help you|Here\'s how|To get started|First, let\'s|Now, let\'s|In conclusion|To summarize)\b',
        r'\b(It is important to note|It should be noted|It is worth mentioning|It is crucial to understand|It is essential to)\b',
        r'\b(Furthermore|Moreover|Additionally|In addition|On the other hand|However|Therefore|Thus|Hence)\b',
        r'\b(As we can see|As mentioned earlier|As discussed|As previously stated|As noted)\b',
        r'\b(Let\'s dive into|Let\'s explore|Let\'s take a look at|Let\'s examine)\b'
    ]
    
    # Claude-specific patterns
    claude_patterns = [
        r'\b(I understand|I can help|I\'d be happy to|I\'ll provide|I can assist)\b',
        r'\b(Here\'s a comprehensive|Here\'s a detailed|Here\'s an in-depth)\b',
        r'\b(To provide context|To give you context|For context|Contextually)\b',
        r'\b(It\'s worth noting|It\'s important to understand|It\'s crucial to recognize)\b'
    ]
    
    # DeepSeek-specific patterns
    deepseek_patterns = [
        r'\b(Let me provide|Let me give you|Let me explain|Let me help)\b',
        r'\b(This is a|This represents|This demonstrates|This shows)\b',
        r'\b(In this case|In this scenario|In this context|In this situation)\b'
    ]
    
    # Generic LLM patterns
    generic_llm_patterns = [
        r'\b(To begin with|To start with|First and foremost|In the first place)\b',
        r'\b(It is widely known|It is commonly accepted|It is generally agreed)\b',
        r'\b(One of the most|One of the key|One of the important|One of the main)\b',
        r'\b(It is worth|It is important|It is crucial|It is essential)\b',
        r'\b(As a result|As a consequence|As an outcome|As a conclusion)\b',
        r'\b(In other words|In simple terms|To put it simply|Simply put)\b',
        r'\b(It should be|It must be|It needs to be|It has to be)\b',
        r'\b(It is clear that|It is obvious that|It is evident that|It is apparent that)\b'
    ]
    
    # Check for ChatGPT patterns
    for pattern in chatgpt_patterns:
        matches = len(re.findall(pattern, text, re.IGNORECASE))
        if matches > 0:
            llm_score += matches * 3
            detected_patterns.append(f"ChatGPT-like pattern: {matches} matches")
    
    # Check for Claude patterns
    for pattern in claude_patterns:
        matches = len(re.findall(pattern, text, re.IGNORECASE))
        if matches > 0:
            llm_score += matches * 2
            detected_patterns.append(f"Claude-like pattern: {matches} matches")
    
    # Check for DeepSeek patterns
    for pattern in deepseek_patterns:
        matches = len(re.findall(pattern, text, re.IGNORECASE))
        if matches > 0:
            llm_score += matches * 2
            detected_patterns.append(f"DeepSeek-like pattern: {matches} matches")
    
    # Check for generic LLM patterns
    for pattern in generic_llm_patterns:
        matches = len(re.findall(pattern, text, re.IGNORECASE))
        if matches > 0:
            llm_score += matches * 1
            detected_patterns.append(f"Generic LLM pattern: {matches} matches")
    
    # Check for overly formal academic language (common in LLMs)
    academic_formality = len(re.findall(r'\b(Furthermore|Moreover|Additionally|Nevertheless|Consequently|Therefore|Thus|Hence|In conclusion|To summarize|It is important to note|It should be noted)\b', text, re.IGNORECASE))
    if academic_formality > 3:
        llm_score += academic_formality * 2
        detected_patterns.append(f"Overly formal academic language: {academic_formality} instances")
    
    # Check for generic opening phrases
    generic_openings = len(re.findall(r'^(Python is|JavaScript is|HTML is|CSS is|SQL is|Machine learning is|Artificial intelligence is|Data science is)', text, re.MULTILINE | re.IGNORECASE))
    if generic_openings > 0:
        llm_score += generic_openings * 5
        detected_patterns.append(f"Generic opening phrases: {generic_openings} instances")
    
    # Check for lack of personal voice
    personal_indicators = len(re.findall(r'\b(I|me|my|myself|we|us|our|ourselves|you|your|yours)\b', text, re.IGNORECASE))
    word_count = len(text.split())
    personal_ratio = personal_indicators / word_count if word_count > 0 else 0
    
    if personal_ratio < 0.01:  # Less than 1% personal pronouns
        llm_score += 10
        detected_patterns.append(f"Very low personal voice: {personal_ratio:.3f} ratio")
    
    # Check for repetitive sentence structures
    sentences = re.split(r'[.!?]+', text)
    if len(sentences) > 5:
        sentence_starters = [s.split()[0].lower() if s.split() else "" for s in sentences[:10]]
        starter_variety = len(set(sentence_starters)) / len(sentence_starters) if sentence_starters else 0
        
        if starter_variety < 0.3:  # Low variety in sentence starters
            llm_score += 8
            detected_patterns.append(f"Repetitive sentence structures: {starter_variety:.3f} variety")
    
    return {
        "llm_score": min(llm_score, 50),  # Cap at 50 points
        "detected_patterns": detected_patterns,
        "pattern_count": len(detected_patterns),
        "confidence": "high" if llm_score > 20 else "medium" if llm_score > 10 else "low"
    }

# Keep existing helper functions for citations, fact-checking, etc.
def extract_citations(text):
    """Extract citations and references from academic text in APA format 7"""
    citations = []
    
    # APA 7 in-text citation patterns
    # Pattern 1: (Author, Year) or (Author et al., Year)
    in_text_pattern1 = r'\(([A-Za-z\s]+(?:et\s+al\.)?,?\s*\d{4})\)'
    in_text_matches1 = re.findall(in_text_pattern1, text)
    
    # Pattern 2: Author (Year) - narrative citations
    in_text_pattern2 = r'([A-Za-z\s]+(?:et\s+al\.)?)\s*\((\d{4})\)'
    in_text_matches2 = re.findall(in_text_pattern2, text)
    
    # Pattern 3: (Author, Year, p. X) - with page numbers
    in_text_pattern3 = r'\(([A-Za-z\s]+(?:et\s+al\.)?,?\s*\d{4},?\s*p\.?\s*\d+)\)'
    in_text_matches3 = re.findall(in_text_pattern3, text)
    
    # Combine all in-text citations
    for match in in_text_matches1:
        citations.append({
            "type": "in_text",
            "text": match.strip(),
            "suspicious": False,
            "format": "APA7"
        })
    
    for author, year in in_text_matches2:
        citations.append({
            "type": "in_text",
            "text": f"{author.strip()}, {year}",
            "suspicious": False,
            "format": "APA7"
        })
    
    for match in in_text_matches3:
        citations.append({
            "type": "in_text",
            "text": match.strip(),
            "suspicious": False,
            "format": "APA7"
        })
    
    # APA 7 reference list patterns
    # Journal article pattern
    journal_pattern = r'^([A-Za-z\s]+(?:et\s+al\.)?,?\s*\(\d{4}\)\.\s*(.+?)\.\s*([A-Za-z\s]+),\s*\d+\(\d+\),\s*\d+[-\d]*\.\s*(?:https?://doi\.org/|DOI:)?([A-Za-z0-9\./]+)?.*?)$'
    journal_matches = re.findall(journal_pattern, text, re.MULTILINE)
    
    for author, title, journal, doi in journal_matches:
        citations.append({
            "type": "reference",
            "text": f"{author.strip()}. {title.strip()}. {journal.strip()}",
            "suspicious": False,
            "format": "APA7",
            "journal": journal.strip(),
            "doi": doi.strip() if doi else None
        })
    
    # Book pattern
    book_pattern = r'^([A-Za-z\s]+(?:et\s+al\.)?,?\s*\(\d{4}\)\.\s*(.+?)\.\s*([A-Za-z\s]+)\.\s*(?:https?://|DOI:)?([A-Za-z0-9\./]+)?.*?)$'
    book_matches = re.findall(book_pattern, text, re.MULTILINE)
    
    for author, title, publisher, url in book_matches:
        citations.append({
            "type": "reference",
            "text": f"{author.strip()}. {title.strip()}. {publisher.strip()}",
            "suspicious": False,
            "format": "APA7",
            "publisher": publisher.strip(),
            "url": url.strip() if url else None
        })
    
    return citations

def verify_journal_existence(journal_name, doi=None):
    """Verify if a journal exists and is legitimate (simplified check)"""
    try:
        # Common legitimate journal indicators
        legitimate_indicators = [
            'journal', 'review', 'proceedings', 'transactions', 'letters',
            'science', 'nature', 'research', 'academic', 'university'
        ]
        
        # Check if journal name contains legitimate indicators
        journal_lower = journal_name.lower()
        has_legitimate_indicators = any(indicator in journal_lower for indicator in legitimate_indicators)
        
        # Check for suspicious patterns (AI-generated journal names)
        suspicious_patterns = [
            'international journal of', 'journal of international',
            'advances in', 'recent advances', 'new developments',
            'comprehensive', 'systematic', 'meta-analysis'
        ]
        
        is_suspicious = any(pattern in journal_lower for pattern in suspicious_patterns)
        
        # DOI validation (basic check)
        doi_valid = False
        if doi:
            doi_valid = doi.startswith('10.') and len(doi) > 10
        
        return {
            "exists": has_legitimate_indicators and not is_suspicious,
            "doi_valid": doi_valid,
            "suspicious": is_suspicious,
            "confidence": "high" if has_legitimate_indicators and doi_valid else "medium"
        }
        
    except Exception as e:
        logger.warning(f"Journal verification failed: {str(e)}")
        return {
            "exists": False,
            "doi_valid": False,
            "suspicious": True,
            "confidence": "low"
        }

def verify_citation_integrity(citations, text):
    """Verify if citations are legitimate or potentially fabricated"""
    integrity_score = 0
    suspicious_citations = []
    verified_journals = []
    
    for citation in citations:
        citation_suspicious = False
        citation_issues = []
        
        # Get citation text for analysis
        citation_text = citation if isinstance(citation, str) else citation.get("text", str(citation))
        
        # Check for common AI citation patterns
        if re.search(r'et\s+al\.', citation_text):
            # AI often overuses "et al."
            integrity_score += 10
            citation_issues.append("Overuse of 'et al.'")
            citation_suspicious = True
        
        # Check for suspicious year patterns
        years = re.findall(r'\d{4}', citation_text)
        if years:
            current_year = 2024
            for year in years:
                year_int = int(year)
                if year_int > current_year or year_int < 1990:
                    suspicious_citations.append(citation_text)
                    integrity_score += 15
                    citation_issues.append(f"Suspicious year: {year}")
                    citation_suspicious = True
        
        # Check for overly generic author names
        if re.search(r'Smith|Johnson|Brown|Davis|Wilson', citation_text, re.IGNORECASE):
            integrity_score += 5
            citation_issues.append("Generic author names")
            citation_suspicious = True
        
        # If citation has journal information, verify it
        if isinstance(citation, dict) and citation.get("journal"):
            journal_verification = verify_journal_existence(
                citation["journal"], 
                citation.get("doi")
            )
            
            if not journal_verification["exists"]:
                integrity_score += 20
                citation_issues.append("Journal verification failed")
                citation_suspicious = True
            
            if journal_verification["suspicious"]:
                integrity_score += 10
                citation_issues.append("Suspicious journal name pattern")
                citation_suspicious = True
            
            verified_journals.append({
                "journal": citation["journal"],
                "verified": journal_verification["exists"],
                "suspicious": journal_verification["suspicious"],
                "doi_valid": journal_verification["doi_valid"]
            })
        
        if citation_suspicious:
            suspicious_citations.append({
                "text": citation if isinstance(citation, str) else citation.get("text", str(citation)),
                "issues": citation_issues
            })
    
    return {
        "integrity_score": min(integrity_score, 100),
        "suspicious_citations": suspicious_citations,
        "total_citations": len(citations),
        "verified_journals": verified_journals,
        "journal_verification_rate": len([j for j in verified_journals if j["verified"]]) / max(len(verified_journals), 1) if verified_journals else 0
    }

# Keep all existing helper functions for document processing and fact-checking...
# [Previous helper functions remain unchanged]

def estimate_tokens(text):
    """
    Estimate token count for text (rough approximation: 1 token ≈ 4 characters for English)
    """
    return len(text) // 4

def create_smart_chunks(text, max_tokens=250000):
    """
    Create intelligent document chunks for Nova Pro analysis.
    Nova Pro limit: 300,000 tokens (leaving 50k for response)
    """
    chunks = []
    
    # If text is small enough, return as single chunk
    if estimate_tokens(text) <= max_tokens:
        return [text]
    
    # Split by sections for academic documents
    academic_sections = re.split(r'(?i)(?:abstract|introduction|methodology|results|discussion|conclusion|references)', text)
    
    if len(academic_sections) > 1:
        # Academic document - chunk by sections
        current_chunk = ""
        for i, section in enumerate(academic_sections):
            if not section.strip():
                continue
                
            # Add section header
            if i > 0:
                section_header = ["Abstract", "Introduction", "Methodology", "Results", "Discussion", "Conclusion", "References"][min(i-1, 6)]
                section = f"\n\n{section_header}\n{section}"
            
            # Check if adding this section would exceed token limit
            test_chunk = current_chunk + section
            if estimate_tokens(test_chunk) > max_tokens and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = section
            else:
                current_chunk = test_chunk
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
    else:
        # Non-academic document - chunk by paragraphs
        paragraphs = text.split('\n\n')
        current_chunk = ""
        
        for paragraph in paragraphs:
            if not paragraph.strip():
                continue
                
            test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            
            if estimate_tokens(test_chunk) > max_tokens and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                current_chunk = test_chunk
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
    
    return chunks

def fact_check_document(text, bedrock):
    """
    Fact-check document content using Nova Pro to detect hallucinations and verify facts.
    Returns detailed analysis of factual accuracy with source citations.
    """
    try:
        # Create smart chunks for fact-checking (smaller chunks for detailed analysis)
        chunks = create_smart_chunks(text, max_tokens=100000)  # Smaller chunks for detailed fact-checking
        logger.info(f"Created {len(chunks)} chunks for fact-checking")
        
        all_fact_checks = []
        
        for i, chunk in enumerate(chunks):
            logger.info(f"Fact-checking chunk {i+1}/{len(chunks)}")
            
            # Use Nova Pro for fact-checking analysis
            fact_check_result = fact_check_single_chunk(chunk, bedrock)
            
            if fact_check_result:
                all_fact_checks.append(fact_check_result)
        
        # Aggregate fact-checking results
        if all_fact_checks:
            return aggregate_fact_check_results(all_fact_checks)
        else:
            return {
                "status": "error",
                "message": "Fact-checking failed for all chunks",
                "factual_accuracy": 50,
                "hallucinations_detected": [],
                "verified_facts": [],
                "misleading_claims": []
            }
            
    except Exception as e:
        logger.error(f"Fact-checking error: {str(e)}")
        return {
            "status": "error",
            "message": f"Fact-checking failed: {str(e)}",
            "factual_accuracy": 50,
            "hallucinations_detected": [],
            "verified_facts": [],
            "misleading_claims": []
        }

def fact_check_single_chunk(chunk_text, bedrock):
    """
    Fact-check a single chunk using Nova Pro with specialized fact-checking prompts.
    """
    try:
        # Fact-checking prompts for Nova Pro
        fact_check_prompts = [
            {
                "instruction": (
                    "Analyze this academic text for factual accuracy and potential hallucinations. "
                    "Identify specific claims, statistics, dates, names, and facts that can be verified. "
                    "Look for: unsupported claims, impossible statistics, anachronistic information, "
                    "contradictory statements, and information that seems too perfect or generic. "
                    "Return JSON with: {\"factual_claims\": [list of specific claims], "
                    "\"potential_hallucinations\": [list of suspicious claims], "
                    "\"verification_needed\": [list of claims needing verification]}"
                ),
                "weight": 0.40
            },
            {
                "instruction": (
                    "Examine this text for misleading information and half-truths. "
                    "Look for: cherry-picked data, misleading statistics, out-of-context quotes, "
                    "oversimplified complex topics, and biased interpretations. "
                    "Return JSON with: {\"misleading_claims\": [list of potentially misleading statements], "
                    "\"bias_indicators\": [list of bias indicators], "
                    "\"context_issues\": [list of context problems]}"
                ),
                "weight": 0.30
            },
            {
                "instruction": (
                    "Verify the credibility of sources and citations mentioned in this text. "
                    "Check for: fake citations, non-existent sources, outdated information, "
                    "misattributed quotes, and circular references. "
                    "Return JSON with: {\"source_issues\": [list of source problems], "
                    "\"citation_accuracy\": [assessment of citation quality], "
                    "\"verification_status\": [overall verification status]}"
                ),
                "weight": 0.30
            }
        ]
        
        fact_check_results = []
        
        for prompt_data in fact_check_prompts:
            try:
                user_content = f"{prompt_data['instruction']}\n\nText to fact-check:\n\n{chunk_text}"
                
                payload = {
                    "messages": [
                        {
                            "role": "user", 
                            "content": [{"text": user_content}]
                        }
                    ],
                    "inferenceConfig": {
                        "maxTokens": 300,  # More tokens for detailed fact-checking
                        "temperature": 0.1,
                        "stopSequences": ["\n\n", "Text to fact-check"]
                    }
                }
                
                response = bedrock.invoke_model(
                    modelId="amazon.nova-pro-v1:0",
                    contentType="application/json",
                    accept="application/json",
                    body=json.dumps(payload)
                )
                
                model_payload = json.loads(response["body"].read())
                raw_output = ""
                
                if "output" in model_payload and "message" in model_payload["output"]:
                    content = model_payload["output"]["message"].get("content", [])
                    if content and isinstance(content, list) and len(content) > 0:
                        raw_output = content[0].get("text", "")
                
                # Parse the fact-checking results
                try:
                    cleaned_output = raw_output.strip()
                    if cleaned_output.startswith('```json'):
                        cleaned_output = cleaned_output.replace('```json', '', 1).strip()
                    if cleaned_output.endswith('```'):
                        cleaned_output = cleaned_output.rstrip('```').strip()
                    
                    parsed = json.loads(cleaned_output)
                    fact_check_results.append({
                        "prompt_type": prompt_data['instruction'][:50] + "...",
                        "weight": prompt_data['weight'],
                        "results": parsed
                    })
                    
                except Exception as e:
                    logger.warning(f"Fact-check parse error: {str(e)}")
                    fact_check_results.append({
                        "prompt_type": prompt_data['instruction'][:50] + "...",
                        "weight": prompt_data['weight'],
                        "results": {"error": str(e)}
                    })
                    
            except Exception as e:
                logger.error(f"Fact-check API error: {str(e)}")
                fact_check_results.append({
                    "prompt_type": prompt_data['instruction'][:50] + "...",
                    "weight": prompt_data['weight'],
                    "results": {"error": str(e)}
                })
        
        # Process and aggregate fact-checking results
        return process_fact_check_results(fact_check_results, chunk_text)
        
    except Exception as e:
        logger.error(f"Single chunk fact-check error: {str(e)}")
        return None

def process_fact_check_results(fact_check_results, chunk_text):
    """
    Process and analyze fact-checking results from Nova Pro.
    """
    try:
        # Extract claims and issues from all prompts
        all_claims = []
        all_hallucinations = []
        all_misleading = []
        all_source_issues = []
        
        for result in fact_check_results:
            if "error" not in result["results"]:
                # Extract factual claims
                if "factual_claims" in result["results"]:
                    all_claims.extend(result["results"]["factual_claims"])
                
                # Extract potential hallucinations
                if "potential_hallucinations" in result["results"]:
                    all_hallucinations.extend(result["results"]["potential_hallucinations"])
                
                # Extract misleading claims
                if "misleading_claims" in result["results"]:
                    all_misleading.extend(result["results"]["misleading_claims"])
                
                # Extract source issues
                if "source_issues" in result["results"]:
                    all_source_issues.extend(result["results"]["source_issues"])
        
        # Calculate factual accuracy score
        total_issues = len(all_hallucinations) + len(all_misleading) + len(all_source_issues)
        total_claims = len(all_claims)
        
        if total_claims > 0:
            factual_accuracy = max(0, 100 - (total_issues / total_claims) * 100)
        else:
            factual_accuracy = 50  # Neutral if no claims detected
        
        return {
            "chunk_text": chunk_text[:200] + "..." if len(chunk_text) > 200 else chunk_text,
            "factual_accuracy": int(factual_accuracy),
            "total_claims": total_claims,
            "total_issues": total_issues,
            "factual_claims": all_claims[:5],  # Limit to first 5 claims
            "hallucinations_detected": all_hallucinations[:5],  # Limit to first 5
            "misleading_claims": all_misleading[:5],  # Limit to first 5
            "source_issues": all_source_issues[:5]  # Limit to first 5
        }
        
    except Exception as e:
        logger.error(f"Process fact-check results error: {str(e)}")
        return {
            "chunk_text": chunk_text[:200] + "..." if len(chunk_text) > 200 else chunk_text,
            "factual_accuracy": 50,
            "total_claims": 0,
            "total_issues": 0,
            "factual_claims": [],
            "hallucinations_detected": [],
            "misleading_claims": [],
            "source_issues": []
        }

def aggregate_fact_check_results(all_fact_checks):
    """
    Aggregate fact-checking results from all chunks.
    """
    try:
        # Calculate overall factual accuracy
        total_accuracy = sum(check["factual_accuracy"] for check in all_fact_checks)
        avg_accuracy = total_accuracy / len(all_fact_checks) if all_fact_checks else 50
        
        # Aggregate all issues
        all_hallucinations = []
        all_misleading = []
        all_source_issues = []
        all_claims = []
        
        for check in all_fact_checks:
            all_hallucinations.extend(check.get("hallucinations_detected", []))
            all_misleading.extend(check.get("misleading_claims", []))
            all_source_issues.extend(check.get("source_issues", []))
            all_claims.extend(check.get("factual_claims", []))
        
        # Remove duplicates and limit results
        unique_hallucinations = list(set(all_hallucinations))[:10]
        unique_misleading = list(set(all_misleading))[:10]
        unique_source_issues = list(set(all_source_issues))[:10]
        unique_claims = list(set(all_claims))[:10]
        
        return {
            "status": "success",
            "factual_accuracy": int(avg_accuracy),
            "total_chunks_analyzed": len(all_fact_checks),
            "total_claims_found": len(unique_claims),
            "total_issues_found": len(unique_hallucinations) + len(unique_misleading) + len(unique_source_issues),
            "hallucinations_detected": unique_hallucinations,
            "misleading_claims": unique_misleading,
            "source_issues": unique_source_issues,
            "verified_facts": unique_claims
        }
        
    except Exception as e:
        logger.error(f"Aggregate fact-check results error: {str(e)}")
        return {
            "status": "error",
            "message": f"Aggregation failed: {str(e)}",
            "factual_accuracy": 50,
            "hallucinations_detected": [],
            "misleading_claims": [],
            "source_issues": [],
            "verified_facts": []
        }

def enhanced_comprehend_analysis(text, comprehend):
    """
    Enhanced Comprehend analysis with smart chunking for large documents.
    """
    try:
        # Comprehend has a 5,000 character limit, so we need to chunk large documents
        if len(text) <= 5000:
            text_for_comprehend = text
        else:
            text_for_comprehend = text[:5000]  # Use first 5000 chars for analysis
        
        # Sentiment analysis
        sentiment_result = comprehend.detect_sentiment(Text=text_for_comprehend, LanguageCode="en")
        sentiment = sentiment_result.get("Sentiment", "NEUTRAL")
        sentiment_scores = sentiment_result.get("SentimentScore", {})
        
        # Key phrases
        key_phrases_result = comprehend.detect_key_phrases(Text=text_for_comprehend, LanguageCode="en")
        key_phrases = [kp["Text"] for kp in key_phrases_result.get("KeyPhrases", [])]
        
        # Named entities
        entities_result = comprehend.detect_entities(Text=text_for_comprehend, LanguageCode="en")
        entities = entities_result.get("Entities", [])
        
        # Dominant language
        language_result = comprehend.detect_dominant_language(Text=text_for_comprehend)
        dominant_language = language_result.get("Languages", [{}])[0].get("LanguageCode", "en")
        
        # Syntax analysis (if available)
        syntax_result = comprehend.detect_syntax(Text=text_for_comprehend, LanguageCode="en")
        syntax_tokens = syntax_result.get("SyntaxTokens", [])
        
        return {
            "sentiment": sentiment,
            "sentiment_scores": sentiment_scores,
            "key_phrases": key_phrases,
            "entities": entities,
            "dominant_language": dominant_language,
            "syntax_tokens": syntax_tokens
        }
        
    except Exception as e:
        logger.warning(f"Enhanced Comprehend analysis failed: {str(e)}")
        return {
            "sentiment": "UNKNOWN",
            "sentiment_scores": {},
            "key_phrases": [],
            "entities": [],
            "dominant_language": "en",
            "syntax_tokens": []
        }

def analyze_text(text, comprehend, bedrock):
    """
    Main analysis function using the new rubric-based scoring system.
    """
    try:
        # Ensure text is a string
        if not isinstance(text, str):
            logger.error(f"Text parameter is not a string: {type(text)}")
            text = str(text) if text else ""
        
        if not text or len(text.strip()) < 10:
            logger.warning("Text is too short or empty")
            return {
                "sentiment": "UNKNOWN",
                "key_phrases": [],
                "ai_score": 50,
                "classification": "Unknown",
                "explanation": "Text too short for analysis",
                "confidence": "low",
                "rubric_breakdown": {},
                "comprehend_analysis": {}
            }
        
        # Enhanced Comprehend analysis
        comprehend_results = enhanced_comprehend_analysis(text, comprehend)
        sentiment = comprehend_results["sentiment"]
        sentiment_scores = comprehend_results["sentiment_scores"]
        key_phrases = comprehend_results["key_phrases"]
        entities = comprehend_results["entities"]
        dominant_language = comprehend_results["dominant_language"]
        syntax_tokens = comprehend_results["syntax_tokens"]

        # Use new rubric-based scoring
        rubric_results = calculate_rubric_score(text, bedrock)
        
        final_score = rubric_results["total_score"]
        classification = rubric_results["classification"]
        confidence = rubric_results["confidence"]
        explanation = rubric_results["explanation"]
        category_breakdown = rubric_results["category_breakdown"]
        
        # Determine confidence level based on analysis quality
        word_count = len(text.split())
        if word_count < 100:
            confidence = "low"
        elif word_count < 300 and confidence == "high":
            confidence = "medium"
        
        # Get fact-checking results from the rubric analysis
        fact_checking_results = None
        if "factuality_hallucination" in category_breakdown:
            fact_category = category_breakdown["factuality_hallucination"]
            fact_checking_results = {
                "status": "success",
                "factual_accuracy": fact_category.get("factual_accuracy", 50),
                "hallucinations_detected": fact_category.get("hallucinations", []),
                "misleading_claims": fact_category.get("misleading_claims", []),
                "source_issues": [],
                "verified_facts": [],
                "source_recommendations": []
            }
        else:
            # Perform dedicated fact-checking if not available in rubric
            try:
                fact_checking_results = fact_check_document(text, bedrock)
            except Exception as e:
                logger.warning(f"Fact-checking failed: {str(e)}")
                fact_checking_results = {
                    "status": "error",
                    "factual_accuracy": 50,
                    "hallucinations_detected": [],
                    "misleading_claims": [],
                    "source_issues": [],
                    "verified_facts": [],
                    "source_recommendations": []
                }
        
        return {
            "sentiment": sentiment,
            "key_phrases": key_phrases,
            "ai_score": final_score,
            "classification": classification,
            "explanation": explanation,
            "confidence": confidence,
            "rubric_breakdown": category_breakdown,
            "comprehend_analysis": {
                "sentiment_scores": sentiment_scores,
                "entities": entities[:10],
                "dominant_language": dominant_language,
                "syntax_complexity": len(syntax_tokens) if syntax_tokens else 0
            },
            "fact_checking": fact_checking_results or {
                "status": "not_available",
                "factual_accuracy": 50,
                "message": "Fact-checking not performed"
            }
        }
        
    except Exception as e:
        logger.error(f"Analysis error: {str(e)}")
        return {
            "sentiment": "UNKNOWN",
            "key_phrases": [],
            "ai_score": 50,
            "classification": "Unknown",
            "explanation": f"Analysis error: {str(e)}",
            "confidence": "low",
            "rubric_breakdown": {},
            "comprehend_analysis": {}
        }

# Keep all existing handler functions unchanged
def lambda_handler(event, context):
    """Main Lambda handler with enhanced CORS support"""
    logger.info(f"Processing event: {json.dumps(event)}")
    
    if event.get("httpMethod") == "OPTIONS":
        return handle_options_request()
    
    try:
        if event.get("body"):
            body = json.loads(event.get("body", "{}"))
            action = body.get("action")
            
            if action == "upload_and_analyze":
                return handle_upload_and_analyze(body)
        
        if event.get("Records"):
            return handle_s3_event(event)
        
        body = json.loads(event.get("body", "{}"))
        bucket_name = body.get("bucket_name", BUCKET_NAME)
        object_key = body.get("object_key", "")
        if not object_key:
            return {
                "statusCode": 400,
                "headers": get_cors_headers(),
                "body": json.dumps({"error": "No object_key provided"})
            }
        
        return analyze_existing_document(bucket_name, object_key)

    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error: {str(e)}")
        return {
            "statusCode": 400,
            "headers": get_cors_headers(),
            "body": json.dumps({"error": "Invalid JSON in request body"})
        }
    except Exception as e:
        logger.error(f"Lambda execution error: {str(e)}")
        return {
            "statusCode": 500,
            "headers": get_cors_headers(),
            "body": json.dumps({"error": str(e)})
        }

def handle_upload_and_analyze(body):
    """Handle file upload and immediate analysis"""
    try:
        file_name = body.get("fileName")
        file_type = body.get("fileType")
        file_content = body.get("fileContent")
        bucket_name = body.get("bucketName", BUCKET_NAME)
        
        if not all([file_name, file_type, file_content]):
            return {
                "statusCode": 400,
                "headers": get_cors_headers(),
                "body": json.dumps({"error": "Missing required fields: fileName, fileType, fileContent"})
            }
        
        s3, dynamodb, comprehend, bedrock = get_clients()
        
        import base64
        file_bytes = base64.b64decode(file_content)
        s3.put_object(
            Bucket=bucket_name,
            Key=file_name,
            Body=file_bytes,
            ContentType=file_type
        )
        
        logger.info(f"Uploaded {file_name} to S3")
        
        return analyze_existing_document(bucket_name, file_name)
        
    except Exception as e:
        logger.error(f"Upload and analyze error: {str(e)}")
        return {
            "statusCode": 500,
            "headers": get_cors_headers(),
            "body": json.dumps({"error": f"Upload failed: {str(e)}"})
        }

def handle_s3_event(event):
    """Handle S3 event triggers"""
    try:
        record = event["Records"][0]
        bucket_name = record["s3"]["bucket"]["name"]
        object_key = unquote_plus(record["s3"]["object"]["key"])
        logger.info(f"Processing S3 event: {bucket_name}/{object_key}")
        
        return analyze_existing_document(bucket_name, object_key)
        
    except Exception as e:
        logger.error(f"S3 event handling error: {str(e)}")
        return {
            "statusCode": 500,
            "headers": get_cors_headers(),
            "body": json.dumps({"error": str(e)})
        }

def analyze_existing_document(bucket_name, object_key):
    """Analyze an existing document in S3"""
    try:
        s3, dynamodb, comprehend, bedrock = get_clients()
        table = dynamodb.Table(TABLE_NAME)

        extracted_text, extraction_method = extract_text_from_document(s3, bucket_name, object_key)
        if len(extracted_text) < 10:
            return {
                "statusCode": 400,
                "headers": get_cors_headers(),
                "body": json.dumps({"error": "Extracted text too short"})
            }

        # Use new rubric-based analysis
        analysis_results = analyze_text(extracted_text, comprehend, bedrock)
        
        ai_score = analysis_results.get("ai_score", 50)
        classification = analysis_results.get("classification", "Unknown")
        explanation = analysis_results.get("explanation", "Analysis completed")
        sentiment = analysis_results.get("sentiment", "NEUTRAL")
        key_phrases = analysis_results.get("key_phrases", [])
        confidence = analysis_results.get("confidence", "medium")
        rubric_breakdown = analysis_results.get("rubric_breakdown", {})
        comprehend_analysis = analysis_results.get("comprehend_analysis", {})

        try:
            obj_metadata = s3.head_object(Bucket=bucket_name, Key=object_key)
            file_size = obj_metadata['ContentLength']
            last_modified = obj_metadata['LastModified'].isoformat()
        except Exception as e:
            logger.warning(f"Metadata fetch failed: {str(e)}")
            file_size, last_modified = 0, ""

        item = {
            "id": str(uuid.uuid4()),
            "timestamp": str(int(time.time())),
            "bucket_name": bucket_name,
            "object_key": object_key,
            "file_size": file_size,
            "last_modified": last_modified,
            "extraction_method": extraction_method,
            "input_text": extracted_text[:1000],
            "text_length": len(extracted_text),
            "ai_score": ai_score,
            "classification": classification,
            "explanation": explanation[:500],
            "sentiment": sentiment,
            "key_phrases": key_phrases[:10],
            "confidence": confidence,
            "analysis_version": "rubric_v1",
            "source_type": "document",
            "rubric_breakdown": rubric_breakdown
        }

        try:
            table.put_item(Item=item)
            logger.info(f"Stored analysis in DynamoDB for {object_key}")
        except Exception as e:
            logger.error(f"Failed DynamoDB write: {str(e)}")

        try:
            result_key = f"analysis-results/{item['id']}.json"
            s3.put_object(
                Bucket=BUCKET_NAME,
                Key=result_key,
                Body=json.dumps(item, indent=2),
                ContentType="application/json"
            )
            logger.info(f"Saved results to s3://{BUCKET_NAME}/{result_key}")
        except Exception as e:
            logger.error(f"Failed S3 write: {str(e)}")

        response_body = {
            "document_id": item["id"],
            "object_key": object_key,
            "ai_score": ai_score,
            "classification": classification,
            "explanation": explanation,
            "sentiment": sentiment,
            "key_phrases": key_phrases,
            "confidence": confidence,
            "text_length": len(extracted_text),
            "extraction_method": extraction_method,
            "input_text": extracted_text[:1000],
            "analysis_version": "rubric_v1",
            "rubric_breakdown": rubric_breakdown,
            "comprehend_analysis": comprehend_analysis,
            "fact_checking": analysis_results.get("fact_checking", {
                "status": "not_available",
                "factual_accuracy": 50,
                "message": "Fact-checking not performed"
            })
        }

        return {
            "statusCode": 200,
            "headers": get_cors_headers(),
            "body": json.dumps(response_body)
        }

    except Exception as e:
        logger.error(f"Document analysis error: {str(e)}")
        return {
            "statusCode": 500,
            "headers": get_cors_headers(),
            "body": json.dumps({"error": str(e)})
        }